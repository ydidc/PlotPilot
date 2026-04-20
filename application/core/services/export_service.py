"""导出服务：生成真实可打开的 DOCX / EPUB / PDF / Markdown。"""
from __future__ import annotations

import html
import io
import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Iterator, List, Optional, Tuple

from domain.novel.repositories.novel_repository import NovelRepository
from domain.novel.repositories.chapter_repository import ChapterRepository
from domain.novel.entities.novel import Novel
from domain.novel.entities.chapter import Chapter
from domain.novel.value_objects.novel_id import NovelId
from domain.novel.value_objects.chapter_id import ChapterId

logger = logging.getLogger(__name__)


def _safe_filename_stem(title: str, max_len: int = 80) -> str:
    t = (title or "novel").strip()
    t = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", t)
    t = t.replace(" ", "_").strip("._") or "novel"
    if len(t) > max_len:
        t = t[:max_len]
    return t


def _novel_id_str(novel: Novel) -> str:
    nid = novel.id
    return nid.value if hasattr(nid, "value") else str(nid)


def _chapter_display_title(ch: Chapter) -> str:
    if ch.title and str(ch.title).strip():
        return str(ch.title).strip()
    return f"第 {ch.number} 章"


def _content_to_html_paragraphs(text: str) -> str:
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    parts: List[str] = []
    for block in raw.split("\n"):
        line = block.strip()
        if line:
            parts.append(f"<p>{html.escape(line)}</p>")
    if not parts:
        return "<p></p>"
    return "\n".join(parts)


def _cjk_font_paths() -> Iterator[Path]:
    env = os.environ.get("PLOTPILOT_EXPORT_CJK_FONT", "").strip()
    if env:
        yield Path(env)
    if os.name == "nt":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        fonts = Path(windir) / "Fonts"
        for name in (
            "msyh.ttf",
            "simhei.ttf",
            "simsun.ttc",
            "msyh.ttc",
            "simkai.ttf",
        ):
            yield fonts / name
    else:
        for p in (
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        ):
            yield Path(p)


class ExportService:
    """导出服务"""

    def __init__(self, novel_repository: NovelRepository, chapter_repository: ChapterRepository):
        self.novel_repository = novel_repository
        self.chapter_repository = chapter_repository

    def export_novel(self, novel_id: str, format: str) -> Tuple[bytes, str, str]:
        try:
            logger.info("开始导出小说: %s, 格式: %s", novel_id, format)
            novel = self.novel_repository.get_by_id(NovelId(novel_id))
            if not novel:
                raise ValueError(f"小说不存在: {novel_id}")
            chapters = self.chapter_repository.list_by_novel(NovelId(novel_id))
            chapters.sort(key=lambda x: x.number)
            logger.info("导出: %s, 章节数 %s", novel.title, len(chapters))
            if format == "epub":
                result = self._export_to_epub(novel, chapters)
            elif format == "pdf":
                result = self._export_to_pdf(novel, chapters)
            elif format == "docx":
                result = self._export_to_docx(novel, chapters)
            elif format == "markdown":
                result = self._export_to_markdown(novel, chapters)
            else:
                raise ValueError(f"不支持的导出格式: {format}")
            logger.info("导出成功，%s 字节", len(result[0]))
            return result
        except ValueError:
            raise
        except Exception as e:
            logger.error("导出小说失败: %s", e, exc_info=True)
            raise

    def export_chapter(self, chapter_id: str, format: str) -> Tuple[bytes, str, str]:
        try:
            logger.info("开始导出章节: %s, 格式: %s", chapter_id, format)
            chapter = self.chapter_repository.get_by_id(ChapterId(chapter_id))
            if not chapter:
                raise ValueError(f"章节不存在: {chapter_id}")
            novel_id = chapter.novel_id.value if hasattr(chapter.novel_id, "value") else chapter.novel_id
            novel = self.novel_repository.get_by_id(NovelId(novel_id))
            if not novel:
                raise ValueError(f"小说不存在: {novel_id}")
            if format == "epub":
                result = self._export_to_epub(novel, [chapter])
            elif format == "pdf":
                result = self._export_to_pdf(novel, [chapter])
            elif format == "docx":
                result = self._export_to_docx(novel, [chapter])
            elif format == "markdown":
                result = self._export_to_markdown(novel, [chapter])
            else:
                raise ValueError(f"不支持的导出格式: {format}")
            data, mime, _ = result
            ext = {"epub": "epub", "pdf": "pdf", "docx": "docx", "markdown": "md"}[format]
            chapter_stem = _safe_filename_stem(
                f"{novel.title or 'novel'}-第{chapter.number}章"
            )
            logger.info("导出成功，%s 字节", len(data))
            return data, mime, f"{chapter_stem}.{ext}"
        except ValueError:
            raise
        except Exception as e:
            logger.error("导出章节失败: %s", e, exc_info=True)
            raise

    def _export_to_epub(self, novel: Novel, chapters: list[Chapter]) -> Tuple[bytes, str, str]:
        from ebooklib import epub

        book = epub.EpubBook()
        uid = _novel_id_str(novel)
        book.set_identifier(f"plotpilot:{uid}")
        book.set_title(novel.title or "未命名")
        book.set_language("zh")
        book.add_author(novel.author or "未知作者")

        intro = epub.EpubHtml(
            title="简介",
            file_name="intro.xhtml",
            lang="zh",
        )
        premise = html.escape((novel.premise or "").strip() or "（无简介）")
        intro.content = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" lang="zh">
<head><title>简介</title><meta charset="utf-8"/></head>
<body>
<h1>{html.escape(novel.title or "未命名")}</h1>
<p>作者：{html.escape(novel.author or "—")}</p>
<p>{premise}</p>
</body>
</html>"""
        book.add_item(intro)

        spine_items: List[epub.EpubHtml] = [intro]
        for i, ch in enumerate(chapters):
            fname = f"chap_{i + 1:03d}.xhtml"
            title_txt = _chapter_display_title(ch)
            title_esc = html.escape(title_txt)
            body = _content_to_html_paragraphs(ch.content or "")
            item = epub.EpubHtml(title=title_txt, file_name=fname, lang="zh")
            item.content = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="zh">
<head><title>{title_esc}</title><meta charset="utf-8"/></head>
<body>
<h1>{title_esc}</h1>
{body}
</body>
</html>"""
            book.add_item(item)
            spine_items.append(item)

        book.toc = tuple([intro] + spine_items[1:])
        book.add_item(epub.EpubNcx())
        # 不使用空 EpubNav（ebooklib 生成 nav 时会解析正文，空文档会触发 lxml Document is empty）
        book.spine = spine_items

        tmp_path: Optional[str] = None
        try:
            fd, tmp_path = tempfile.mkstemp(suffix=".epub")
            os.close(fd)
            epub.write_epub(tmp_path, book, {})
            with open(tmp_path, "rb") as f:
                data = f.read()
        finally:
            if tmp_path and os.path.isfile(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        stem = _safe_filename_stem(novel.title)
        return data, "application/epub+zip", f"{stem}.epub"

    def _try_register_cjk_font(self, pdf) -> bool:
        for path in _cjk_font_paths():
            if not path.is_file():
                continue
            try:
                pdf.add_font("PlotExportCJK", "", str(path), uni=True)
                return True
            except Exception as e:
                logger.debug("PDF 跳过字体 %s: %s", path, e)
        return False

    def _export_to_pdf(self, novel: Novel, chapters: list[Chapter]) -> Tuple[bytes, str, str]:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=14)
        font = "Helvetica"
        if self._try_register_cjk_font(pdf):
            font = "PlotExportCJK"

        def add_text(size: float, text: str, line_h: float) -> None:
            pdf.set_font(font, size=size)
            body = (text or "").strip() or " "
            try:
                pdf.multi_cell(0, line_h, body, new_x="LMARGIN", new_y="NEXT")
            except Exception as e:
                logger.warning("PDF multi_cell 回退: %s", e)
                pdf.set_font("Helvetica", size=size)
                safe = (text or "").encode("ascii", errors="replace").decode("ascii")
                pdf.multi_cell(0, line_h, safe or " ", new_x="LMARGIN", new_y="NEXT")

        pdf.add_page()
        add_text(16, novel.title or "未命名", 9)
        pdf.ln(2)
        add_text(
            11,
            f"作者：{novel.author or '—'}\n简介：{(novel.premise or '').strip() or '—'}",
            6,
        )
        pdf.ln(4)

        for ch in chapters:
            add_text(14, _chapter_display_title(ch), 8)
            pdf.ln(1)
            add_text(11, (ch.content or "").strip() or "（无正文）", 6)
            pdf.ln(6)

        out = pdf.output()
        if isinstance(out, str):
            data = out.encode("latin-1")
        elif isinstance(out, bytearray):
            data = bytes(out)
        else:
            data = out
        stem = _safe_filename_stem(novel.title)
        return data, "application/pdf", f"{stem}.pdf"

    def _export_to_docx(self, novel: Novel, chapters: list[Chapter]) -> Tuple[bytes, str, str]:
        from docx import Document

        doc = Document()
        doc.add_heading(novel.title or "未命名", level=0)
        doc.add_paragraph(f"作者：{novel.author or '—'}")
        p_pre = doc.add_paragraph()
        p_pre.add_run("简介：").bold = True
        p_pre.add_run((novel.premise or "").strip() or "（无）")

        for ch in chapters:
            doc.add_heading(_chapter_display_title(ch), level=1)
            content = ch.content or ""
            if not content.strip():
                doc.add_paragraph("（无正文）")
                continue
            for line in content.splitlines():
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        stem = _safe_filename_stem(novel.title)
        return (
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{stem}.docx",
        )

    def _export_to_markdown(self, novel: Novel, chapters: list[Chapter]) -> Tuple[bytes, str, str]:
        lines: List[str] = [
            f"# {novel.title or '未命名'}",
            "",
            f"**作者**: {novel.author or '—'}",
            "",
            "## 简介",
            "",
            (novel.premise or "").strip() or "（无）",
            "",
        ]
        for ch in chapters:
            lines.append(f"## {_chapter_display_title(ch)}")
            lines.append("")
            lines.append((ch.content or "").strip() or "（无正文）")
            lines.append("")
        text = "\n".join(lines)
        stem = _safe_filename_stem(novel.title)
        return text.encode("utf-8"), "text/markdown; charset=utf-8", f"{stem}.md"
